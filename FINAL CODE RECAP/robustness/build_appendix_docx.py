#!/usr/bin/env python3
"""
build_appendix_docx.py
======================
Builds capstone_robustness_appendix.docx from pre-generated PNG figures.

Run from:  /Users/leoss/Desktop/GitHub/Capstone/FINAL CODE RECAP
    python robustness/build_appendix_docx.py

Output:  Final/capstone_robustness_appendix.docx
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

APPENDIX_DIR = os.path.join("Final", "appendix")
OUT_PATH     = os.path.join("Final", "capstone_robustness_appendix.docx")

# Re-save all PNGs at 600 DPI using matplotlib
import matplotlib.pyplot as plt
import matplotlib.image as mpimg

def hires_path(name):
    """Return path to a hi-res version of figure (resampled to 600 DPI for docx)."""
    src = os.path.join(APPENDIX_DIR, f"{name}.png")
    dst = os.path.join(APPENDIX_DIR, f"{name}_600.png")
    if not os.path.exists(dst) or os.path.getmtime(src) > os.path.getmtime(dst):
        img = mpimg.imread(src)
        h, w = img.shape[:2]
        fig = plt.figure(figsize=(w / 150, h / 150), dpi=150)
        ax  = fig.add_axes([0, 0, 1, 1])
        ax.imshow(img)
        ax.axis('off')
        fig.savefig(dst, dpi=600, bbox_inches='tight', pad_inches=0)
        plt.close(fig)
    return dst


def set_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
    return p


def add_fig(doc, png_path, caption, width_in=6.5):
    if not os.path.exists(png_path):
        doc.add_paragraph(f"[MISSING: {png_path}]")
        return
    doc.add_picture(png_path, width=Inches(width_in))
    # Centre the picture
    last = doc.paragraphs[-1]
    last.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Caption
    cap = doc.add_paragraph(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.runs[0]
    cap_run.font.size = Pt(9)
    cap_run.font.italic = True
    cap_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    doc.add_paragraph()


# ── Build document ─────────────────────────────────────────────────────────────

doc = Document()

# Page margins
from docx.shared import Mm
section = doc.sections[0]
section.top_margin    = Mm(20)
section.bottom_margin = Mm(20)
section.left_margin   = Mm(25)
section.right_margin  = Mm(25)

# Default font
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)

# ── COVER ──────────────────────────────────────────────────────────────────────

doc.add_heading('Robustness Appendix', level=0)
intro = doc.add_paragraph(
    'This appendix documents two independent robustness checks applied to the '
    'empirical models in the main paper: (I) a parametric bootstrap with full '
    're-imputation (B = 200) and (II) leave-one-country-out (LOCO) analysis '
    'over the 54-country sample. '
    'Bootstrap resamples countries with replacement; each draw undergoes the '
    'full NB3 imputation pipeline before model estimation. '
    'LOCO refits all models on the 53-country subsample obtained by removing '
    'each country in turn.'
)
intro.paragraph_format.space_after = Pt(12)
doc.add_paragraph()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION I — BOOTSTRAP ROBUSTNESS
# ══════════════════════════════════════════════════════════════════════════════

set_heading(doc, 'I. Bootstrap Robustness (B = 200)', level=1)

doc.add_paragraph(
    'Countries are resampled with replacement (N = 54 draws per iteration). '
    'Each bootstrap dataset undergoes linear interpolation followed by '
    'k-NN imputation (k = 5) before model estimation, mirroring the original '
    'data preparation pipeline. '
    'Confidence intervals are the 2.5th–97.5th percentiles of the bootstrap distribution.'
).paragraph_format.space_after = Pt(6)

set_heading(doc, 'I.1  OLS Regression — Coefficient Stability', level=2)
doc.add_paragraph(
    'Figure I.1 plots 95% bootstrap confidence intervals for each OLS coefficient. '
    'Green bars indicate the CI excludes zero (significant); grey bars include zero. '
    'Tick marks show the point estimate from the original full-sample model '
    '(clustered standard errors). '
    'Both Model 3a (without lagged ECI) and Model 3b (with lagged ECI) are shown.'
).paragraph_format.space_after = Pt(4)

add_fig(doc, hires_path('fig_A1_boot_reg_coefs'),
        'Figure I.1 — Bootstrap coefficient 95% CIs, OLS regression Models 3a and 3b. '
        'Green = CI excludes zero; grey = CI includes zero. Tick = original estimate.')

set_heading(doc, 'I.2  ML Models — Test R² Distribution', level=2)
doc.add_paragraph(
    'Figure I.2 shows the bootstrap distribution of out-of-sample test R² '
    'for each ML model (LASSO, Ridge, Elastic Net, Random Forest). '
    'Shaded violins show the full distribution; thick bars show the IQR; '
    'white dots show the median. Annotations report the median and 95% CI.'
).paragraph_format.space_after = Pt(4)

add_fig(doc, hires_path('fig_A2_boot_ml_r2'),
        'Figure I.2 — Bootstrap distribution of test R² for ML models (B = 200). '
        'Thick bar = IQR; whisker = 95% CI; dot = median.')

set_heading(doc, 'I.3  Random Forest — Feature Importance', level=2)
doc.add_paragraph(
    'Figure I.3 presents bootstrap 95% CIs for Random Forest feature importances '
    '(Gini criterion), showing the top 12 features ranked by median importance.'
).paragraph_format.space_after = Pt(4)

add_fig(doc, hires_path('fig_A3_boot_rf_importance'),
        'Figure I.3 — Bootstrap 95% CIs for RF Gini feature importance, top 12 features.')

set_heading(doc, 'I.4  LASSO — Selection Frequency', level=2)
doc.add_paragraph(
    'Figure I.4 shows the proportion of bootstrap iterations in which each '
    'feature receives a non-zero LASSO coefficient. '
    'Features above the 90% threshold are highlighted in green.'
).paragraph_format.space_after = Pt(4)

add_fig(doc, hires_path('fig_A4_boot_lasso_selection'),
        'Figure I.4 — LASSO selection frequency across bootstrap iterations. '
        'Green = selected in >90% of runs; blue = >50%; grey = ≤50%.')

set_heading(doc, 'I.5  Penalised ML Models — Coefficient Stability', level=2)
doc.add_paragraph(
    'Figure I.5 plots 95% bootstrap CIs for standardised coefficients in the '
    'three penalised ML models (LASSO, Elastic Net, Ridge), sorted by absolute '
    'median value. Green indicates significance (CI excludes zero).'
).paragraph_format.space_after = Pt(4)

add_fig(doc, hires_path('fig_A4b_boot_ml_coefs'),
        'Figure I.5 — Bootstrap 95% CIs for penalised ML coefficients (B = 200). '
        'Left: LASSO. Centre: Elastic Net. Right: Ridge.')

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION II — ML LOCO
# ══════════════════════════════════════════════════════════════════════════════

set_heading(doc, 'II. LOCO Robustness — ML Models', level=1)

doc.add_paragraph(
    'Each of the 54 countries is excluded in turn; all four ML models are '
    're-estimated on the remaining 53-country sample and evaluated on the '
    'held-out test period (2015–2019). '
    'Coefficient/importance stability is assessed relative to the full-sample model.'
).paragraph_format.space_after = Pt(6)

set_heading(doc, 'II.1  Coefficient Stability — LASSO, Elastic Net, Ridge', level=2)
doc.add_paragraph(
    'Figure II.1 shows the min-max range of standardised coefficients across the '
    '54 LOCO runs for each penalised model. '
    'Dots show the LOCO mean; bars show the full range.'
).paragraph_format.space_after = Pt(4)

add_fig(doc, hires_path('fig_A5_loco_ml_coef'),
        'Figure II.1 — LOCO coefficient stability for penalised ML models '
        '(N = 54 LOO runs). Dot = LOCO mean; bar = min-max range.')

set_heading(doc, 'II.2  Test R² Stability', level=2)
doc.add_paragraph(
    'Figure II.2 shows the out-of-sample test R² for each excluded country '
    'and for all four ML models. Countries are sorted by mean R² across models. '
    'Narrow spread across the 54 bars indicates that no single country is disproportionately '
    'influential.'
).paragraph_format.space_after = Pt(4)

add_fig(doc, hires_path('fig_A6_loco_ml_r2'),
        'Figure II.2 — LOCO test R² per excluded country, all four ML models. '
        'Countries sorted by mean R² across models.')

doc.add_page_break()


# ══════════════════════════════════════════════════════════════════════════════
# SECTION III — REGRESSION LOCO
# ══════════════════════════════════════════════════════════════════════════════

set_heading(doc, 'III. LOCO Robustness — OLS Regression', level=1)

doc.add_paragraph(
    'The OLS regression LOCO repeats the exercise for Models 3a and 3b. '
    'Plain OLS is used within each LOO iteration to track coefficient stability; '
    'inference in the main paper uses country-clustered standard errors. '
    'The specification includes seven regressors and four interaction terms '
    '(HCI × Production, GFCF × Production, HCI × Forestry Rents, '
    'GFCF × Forestry Rents), matching the NB6 specification exactly.'
).paragraph_format.space_after = Pt(6)

set_heading(doc, 'III.1  Coefficient Stability', level=2)
doc.add_paragraph(
    'Figure III.1 shows the 5th–95th percentile range (thick bar) and full '
    'min-max range (thin bar) of OLS coefficients across the 54 LOCO runs. '
    'Dots show the LOCO mean.'
).paragraph_format.space_after = Pt(4)

add_fig(doc, hires_path('fig_A7_loco_reg_coef'),
        'Figure III.1 — LOCO coefficient stability for OLS regression Models 3a and 3b '
        '(N = 54 LOO runs). Thick bar = 5–95th pct; thin bar = min-max; dot = LOCO mean.')

set_heading(doc, 'III.2  R² Stability', level=2)
doc.add_paragraph(
    'Figure III.2 shows R² per excluded country for Model 3a (top) and Model 3b '
    '(bottom), sorted independently within each panel. '
    'Dashed line shows the LOCO mean R².'
).paragraph_format.space_after = Pt(4)

add_fig(doc, hires_path('fig_A8_loco_reg_r2'),
        'Figure III.2 — LOCO R² per excluded country, Models 3a and 3b separately. '
        'Dashed line = LOCO mean. Countries sorted by R².')

# ── Save ──────────────────────────────────────────────────────────────────────

doc.save(OUT_PATH)
print(f"\nSaved: {OUT_PATH}")
